from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from django.http import HttpResponse
from django.utils.timezone import localtime
import re

# --- COLORES (Coinciden con el tema web) ---
COLOR_PRIMARY = RGBColor(79, 70, 229)    # Indigo
COLOR_SECONDARY = RGBColor(16, 185, 129) # Emerald
COLOR_TEXT_MAIN = RGBColor(31, 41, 55)   # Gray 800
COLOR_TEXT_MUTED = RGBColor(107, 114, 128) # Gray 500
COLOR_BG_HEADER = "EEF2FF" # Indigo 50 (Hex string for XML)

def _v(x, dash="—"):
    s = "" if x is None else str(x).strip()
    return s if s else dash

def _set_base(doc):
    # Márgenes
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin  = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    # Estilo Normal
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI' # Fuente más moderna
    font.size = Pt(10)
    font.color.rgb = COLOR_TEXT_MAIN
    
    # Estilo Heading 1 (Título del documento)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Segoe UI'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_PRIMARY
    h1.paragraph_format.space_after = Pt(12)
    
    # Estilo Heading 2 (Títulos de sección)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Segoe UI'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_PRIMARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)
    # Borde inferior para H2
    pPr = h2._element.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E5E7EB')
    pbdr.append(bottom)
    pPr.append(pbdr)

def _fill(cell, hex_color):
    """Relleno de celda."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), hex_color)

def _create_header(doc, title, subtitle=None):
    p = doc.add_paragraph(style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)
    
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_TEXT_MUTED
        p_sub.paragraph_format.space_after = Pt(24)

def _kv_table(doc, pairs, col1_width=5.0, col2_width=11.0):
    """Tabla clave-valor limpia sin bordes verticales internos."""
    t = doc.add_table(rows=0, cols=2)
    t.style = 'Table Grid' # Usamos grid base y lo modificamos si es necesario
    t.autofit = False
    
    for k, v in pairs:
        row = t.add_row()
        c0, c1 = row.cells
        
        # Clave
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(str(k))
        r0.bold = True
        r0.font.color.rgb = COLOR_TEXT_MAIN
        _fill(c0, "F9FAFB") # Gray 50
        
        # Valor
        p1 = c1.paragraphs[0]
        p1.add_run(_v(v))
        
        # Anchos
        c0.width = Cm(col1_width)
        c1.width = Cm(col2_width)
        
    doc.add_paragraph() # Espacio después de la tabla

def _block_content(doc, title, text):
    """Bloque de contenido con título pequeño."""
    if not text or not text.strip():
        return # No mostrar bloques vacíos
        
    p_title = doc.add_paragraph()
    r_title = p_title.add_run(title)
    r_title.bold = True
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = COLOR_TEXT_MUTED
    p_title.paragraph_format.space_after = Pt(2)
    
    p_text = doc.add_paragraph(_v(text))
    p_text.paragraph_format.space_after = Pt(12)

def generar_docx_informe(informe):
    est = informe.estudiante
    curso_obj = getattr(est, "curso", None)
    estb = getattr(curso_obj, "establecimiento", None) if curso_obj else None

    nombre_est = informe.nombre_establecimiento or (getattr(estb, "nombre", "") or "")
    rbd        = informe.rbd or (getattr(estb, "rbd", "") or "")
    dep_disp   = getattr(informe, "get_dependencia_display", None)
    dependencia_val = dep_disp() if dep_disp and getattr(informe, "dependencia", None) else (
        getattr(estb, "get_dependencia_display", None)() if estb and hasattr(estb, "dependencia") else ""
    )
    curso_str  = informe.curso or (str(curso_obj) if curso_obj else "")
    edad_str   = str(informe.edad) if informe.edad is not None else ""

    doc = Document()
    _set_base(doc)

    # --- PORTADA ---
    _create_header(doc, "INFORME PIE", nombre_est)

    # --- I. IDENTIFICACIÓN ---
    doc.add_paragraph("I. Identificación", style='Heading 2')
    
    # Usamos una tabla invisible para layout de 2 columnas si fuera necesario, 
    # pero para KV simple, una tabla unificada se ve mejor.
    
    data_ident = [
        ("Estudiante", f"{est.nombres} {est.apellidos}"),
        ("RUN", est.run),
        ("Curso", curso_str),
        ("Edad", edad_str),
        ("Establecimiento", nombre_est),
        ("RBD", rbd),
        ("Dependencia", dependencia_val),
        ("Diagnóstico", informe.diagnostico),
        ("Período", f"{_v(getattr(informe,'periodo_inicio',''))} al {_v(getattr(informe,'periodo_fin',''))}")
    ]
    _kv_table(doc, data_ident)

    # --- II. PLANIFICACIÓN ---
    doc.add_paragraph("II. Planificación y Apoyos", style='Heading 2')
    
    data_plan = [
        ("Objetivos Generales", informe.objetivos_generales),
        ("Estrategias de Apoyo", informe.estrategias_apoyo),
        ("Recursos Utilizados", informe.recursos_utilizados),
    ]
    if informe.frecuencia_apoyo:
        data_plan.append(("Frecuencia de Apoyo", informe.frecuencia_apoyo))
        
    _kv_table(doc, data_plan)

    # --- III. EVALUACIÓN ---
    doc.add_paragraph("III. Evaluación y Seguimiento", style='Heading 2')
    
    data_eval = [
        ("Logros Alcanzados", informe.logros_alcanzados),
        ("Dificultades Detectadas", informe.dificultades_detectadas),
        ("Sugerencias", informe.sugerencias),
        ("Antecedentes Relevantes", informe.antecedentes),
        ("Evaluación General", informe.evaluacion),
    ]
    _kv_table(doc, data_eval)

    # --- IV. OBSERVACIONES ---
    doc.add_paragraph("IV. Observaciones Finales", style='Heading 2')
    _kv_table(doc, [("Observaciones", informe.observaciones)])

    # --- V. RESPONSABLE ---
    doc.add_paragraph("V. Profesional Responsable", style='Heading 2')
    _kv_table(doc, [
        ("Nombre", _v(getattr(informe.profesional, "nombre_completo", "No asignado"))),
        ("RUT", _v(getattr(informe, "rut_profesional", ""))),
        ("Fecha Emisión", localtime(informe.fecha_creacion).strftime("%d-%m-%Y")),
    ])

    # --- VI. COLABORADORES ---
    cols = (
        informe.colaboradores.select_related("profesional")
        .prefetch_related("actividades")
        .order_by("profesional__nombre_completo")
    )

    if cols.exists():
        doc.add_page_break()
        doc.add_paragraph("VI. Registro de Colaboración", style='Heading 2')
        
        for col in cols:
            # Encabezado del colaborador
            p_col = doc.add_paragraph()
            r_col = p_col.add_run(f"{col.profesional.nombre_completo}")
            r_col.bold = True
            r_col.font.size = Pt(12)
            r_col.font.color.rgb = COLOR_SECONDARY
            
            p_meta = doc.add_paragraph()
            p_meta.add_run(f"{col.profesional.especialidad} • {col.get_rol_display()}").font.color.rgb = COLOR_TEXT_MUTED
            
            if col.resumen_aporte:
                _kv_table(doc, [("Resumen del Aporte", col.resumen_aporte)])

            # Tabla de actividades
            acts = list(col.actividades.all().order_by("fecha", "id"))
            if acts:
                t = doc.add_table(rows=1, cols=4)
                t.style = 'Table Grid'
                t.autofit = False
                
                # Headers
                headers = ["Fecha", "Tipo", "Objetivo / Actividad", "Observaciones"]
                widths = [2.5, 3.0, 6.0, 4.5]
                
                for i, h in enumerate(headers):
                    cell = t.rows[0].cells[i]
                    cell.text = h
                    _fill(cell, COLOR_BG_HEADER)
                    cell.paragraphs[0].runs[0].bold = True
                    cell.width = Cm(widths[i])
                
                for a in acts:
                    row = t.add_row()
                    row.cells[0].text = a.fecha.strftime("%d/%m")
                    row.cells[0].width = Cm(widths[0])
                    
                    tipo_txt = a.get_tipo_display() if callable(getattr(a, "get_tipo_display", None)) else str(a.tipo)
                    row.cells[1].text = tipo_txt
                    row.cells[1].width = Cm(widths[1])
                    
                    row.cells[2].text = a.objetivo or ""
                    row.cells[2].width = Cm(widths[2])
                    
                    row.cells[3].text = a.observaciones or ""
                    row.cells[3].width = Cm(widths[3])
                
                doc.add_paragraph() # Espacio
            
            doc.add_paragraph() # Espacio entre colaboradores

    return doc
